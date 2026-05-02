import io
import zipfile
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
from reportlab.lib import colors
from PIL import Image as PILImage


def _extract_images_from_docx(file_bytes: bytes) -> dict:
    """
    Returns a dict mapping relationship id → PIL Image object.
    Word stores images inside the zip at word/media/
    and links them via word/_rels/document.xml.rels
    """
    images = {}
    with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
        # Parse relationships to map rId → filename
        rel_map = {}
        try:
            with z.open("word/_rels/document.xml.rels") as f:
                import xml.etree.ElementTree as ET
                tree = ET.parse(f)
                for rel in tree.getroot():
                    rid = rel.attrib.get("Id")
                    target = rel.attrib.get("Target", "")
                    if "media/" in target:
                        rel_map[rid] = "word/" + target.lstrip("/")
        except KeyError:
            return images

        # Load the actual image bytes for each rel
        for rid, path in rel_map.items():
            try:
                with z.open(path) as img_file:
                    img_bytes = img_file.read()
                    pil_img = PILImage.open(io.BytesIO(img_bytes))
                    if pil_img.mode in ("RGBA", "P", "LA"):
                        pil_img = pil_img.convert("RGB")
                    images[rid] = pil_img
            except Exception:
                continue
    return images


def _para_has_image(para) -> list:
    """Return list of rId strings for any inline images in this paragraph."""
    rids = []
    for run in para.runs:
        # Inline images: w:drawing > wp:inline > a:graphic > ... > r:embed
        for drawing in run._element.findall(f".//{qn('a:blip')}"):
            rid = drawing.attrib.get(f"{{{drawing.nsmap.get('r', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships')}}}embed")
            if rid:
                rids.append(rid)
        # Also check direct r:embed on blip
        for blip in run._element.iter():
            for attr_key, attr_val in blip.attrib.items():
                if "embed" in attr_key and attr_val.startswith("rId"):
                    if attr_val not in rids:
                        rids.append(attr_val)
    return rids


def _pil_to_rl_image(pil_img: PILImage.Image, max_width: float, max_height: float) -> RLImage:
    """Convert a PIL image to a ReportLab Image, scaled to fit the page."""
    buf = io.BytesIO()
    pil_img.save(buf, format="JPEG", quality=88)
    buf.seek(0)

    orig_w, orig_h = pil_img.size
    aspect = orig_h / orig_w

    draw_w = min(max_width, orig_w)
    draw_h = draw_w * aspect
    if draw_h > max_height:
        draw_h = max_height
        draw_w = draw_h / aspect

    return RLImage(buf, width=draw_w, height=draw_h)


def docx_to_pdf(file_bytes: bytes) -> bytes:
    doc = Document(io.BytesIO(file_bytes))
    buffer = io.BytesIO()

    page_w, page_h = A4
    margin = 20 * mm
    usable_w = page_w - 2 * margin
    usable_h = page_h - 2 * margin

    # Extract all embedded images upfront
    embedded_images = _extract_images_from_docx(file_bytes)

    pdf = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=margin,
        leftMargin=margin,
        topMargin=margin,
        bottomMargin=margin,
    )

    styles = getSampleStyleSheet()
    normal_style = ParagraphStyle(
        "custom_normal",
        parent=styles["Normal"],
        fontSize=11,
        leading=16,
        spaceAfter=6,
    )
    heading1_style = ParagraphStyle(
        "custom_h1",
        parent=styles["Heading1"],
        fontSize=14,
        leading=20,
        spaceAfter=10,
        spaceBefore=12,
        textColor=colors.HexColor("#1a1a2e"),
    )
    heading2_style = ParagraphStyle(
        "custom_h2",
        parent=styles["Heading2"],
        fontSize=12,
        leading=18,
        spaceAfter=8,
        spaceBefore=8,
    )

    story = []

    for para in doc.paragraphs:
        # Check for embedded images in this paragraph
        image_rids = _para_has_image(para)
        for rid in image_rids:
            if rid in embedded_images:
                rl_img = _pil_to_rl_image(embedded_images[rid], usable_w, usable_h * 0.8)
                story.append(Spacer(1, 6))
                story.append(rl_img)
                story.append(Spacer(1, 6))

        # Then handle the text of this paragraph
        text = para.text.strip()
        if not text:
            if not image_rids:           # avoid double-spacing after an image
                story.append(Spacer(1, 6))
            continue

        style_name = para.style.name.lower()
        if "heading 1" in style_name:
            story.append(Paragraph(text, heading1_style))
        elif "heading" in style_name:
            story.append(Paragraph(text, heading2_style))
        else:
            story.append(Paragraph(text, normal_style))

    # Tables
    for table in doc.tables:
        data = []
        for row in table.rows:
            data.append([cell.text for cell in row.cells])
        if data:
            t = Table(data, repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4a4e69")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f2f2f2")]),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("PADDING", (0, 0), (-1, -1), 6),
            ]))
            story.append(Spacer(1, 8))
            story.append(t)

    if not story:
        story.append(Paragraph("(empty document)", normal_style))

    pdf.build(story)
    buffer.seek(0)
    return buffer.read()


def image_to_pdf(file_bytes: bytes, filename: str) -> bytes:
    """Convert a standalone image (jpg/png/webp) to a PDF page."""
    buffer = io.BytesIO()
    img = PILImage.open(io.BytesIO(file_bytes))

    try:
        from PIL import ImageOps
        img = ImageOps.exif_transpose(img)
    except Exception:
        pass

    if img.mode in ("RGBA", "P", "LA"):
        img = img.convert("RGB")

    img_buffer = io.BytesIO()
    img.save(img_buffer, format="JPEG", quality=90)
    img_buffer.seek(0)

    page_w, page_h = A4
    margin = 20 * mm
    draw_w = page_w - 2 * margin
    draw_h = draw_w * (img.height / img.width)

    if draw_h > page_h - 2 * margin:
        draw_h = page_h - 2 * margin
        draw_w = draw_h * (img.width / img.height)

    pdf = SimpleDocTemplate(buffer, pagesize=A4,
                            rightMargin=margin, leftMargin=margin,
                            topMargin=margin, bottomMargin=margin)
    pdf.build([RLImage(img_buffer, width=draw_w, height=draw_h)])
    buffer.seek(0)
    return buffer.read()